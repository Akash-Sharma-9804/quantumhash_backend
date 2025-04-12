

# import sys
# import json
# import pytesseract
# from PIL import Image
# from io import BytesIO
# from pdfminer.high_level import extract_pages
# from pdfminer.layout import LTTextContainer
# from docx import Document
# import pandas as pd
# import base64
# import tempfile
# import fitz  # PyMuPDF


# def extract_docx_text(content):
#     with BytesIO(content) as f:
#         doc = Document(f)
#         return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


# def extract_txt_text(content):
#     return content.decode('utf-8').strip()


# def extract_image_text(content):
#     with BytesIO(content) as f:
#         img = Image.open(f).convert("RGB")
#         text = pytesseract.image_to_string(img)
#         return text.strip()


# def extract_excel_text(content):
#     with BytesIO(content) as f:
#         xls = pd.read_excel(f, sheet_name=None)
#         output = ""
#         for sheet_name, df in xls.items():
#             output += f"\n--- Sheet: {sheet_name} ---\n"
#             output += df.to_string(index=False)
#         return output.strip()


# def extract_pdf_pagewise(content):
#     text_by_page = []

#     # Step 1: Try pdfminer.six
#     try:
#         with BytesIO(content) as f:
#             for i, page_layout in enumerate(extract_pages(f), start=1):
#                 page_text = ""
#                 for element in page_layout:
#                     if isinstance(element, LTTextContainer):
#                         page_text += element.get_text()
#                 text_by_page.append(f"\n--- Page {i} ---\n{page_text.strip() or '[No text found]'}")
#     except Exception as e:
#         print(f"⚠️ PDFMiner failed: {str(e)}", file=sys.stderr)

#     # Step 2: Fallback to OCR if too short or mostly empty
#     if not text_by_page or sum(len(p) for p in text_by_page) < 500:
#         text_by_page = []
#         try:
#             with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
#                 tmp.write(content)
#                 tmp.flush()
#                 doc = fitz.open(tmp.name)
#                 for i, page in enumerate(doc, start=1):
#                     text = page.get_text()
#                     if text.strip():
#                         text_by_page.append(f"\n--- Page {i} ---\n{text.strip()}")
#                     else:
#                         pix = page.get_pixmap(dpi=300)
#                         img = Image.open(BytesIO(pix.tobytes("png"))).convert("RGB")
#                         ocr_text = pytesseract.image_to_string(img)
#                         text_by_page.append(f"\n--- OCR Page {i} ---\n{ocr_text.strip() or '[No text found]'}")
#         except Exception as e:
#             return f"[OCR fallback failed: {str(e)}]"

#     return "\n".join(text_by_page).strip()


# def main():
#     try:
#         input_data = json.loads(sys.stdin.read())
#         file_content = base64.b64decode(input_data["buffer"])
#         mime_type = input_data["mimeType"]

#         if mime_type == "application/pdf":
#             result = extract_pdf_pagewise(file_content)
#         elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             result = extract_docx_text(file_content)
#         elif mime_type.startswith("image"):
#             result = extract_image_text(file_content)
#         elif mime_type == "text/plain":
#             result = extract_txt_text(file_content)
#         elif "spreadsheet" in mime_type or "excel" in mime_type:
#             result = extract_excel_text(file_content)
#         else:
#             result = "[Unsupported file type]"

#         # ✅ Only send valid JSON to stdout
#         print(json.dumps({"text": result}))

#     except Exception as e:
#         print(json.dumps({"error": str(e)}))


# if __name__ == "__main__":
#     main()



import sys
import json
import pytesseract
from PIL import Image
from io import BytesIO
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer
from docx import Document
import pandas as pd
import base64
import tempfile
import fitz  # PyMuPDF


def extract_docx_text(content):
    with BytesIO(content) as f:
        doc = Document(f)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


def extract_txt_text(content):
    return content.decode('utf-8').strip()


def extract_image_text(content):
    with BytesIO(content) as f:
        img = Image.open(f).convert("RGB")
        text = pytesseract.image_to_string(img)
        return text.strip()


def extract_excel_text(content):
    with BytesIO(content) as f:
        xls = pd.read_excel(f, sheet_name=None)
        output = ""
        for sheet_name, df in xls.items():
            output += f"\n--- Sheet: {sheet_name} ---\n"
            output += df.to_string(index=False)
        return output.strip()


def extract_pdf_pagewise(content):
    text_by_page = []

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(content)
        tmp.flush()
        doc = fitz.open(tmp.name)

        for i, page in enumerate(doc, start=1):
            # --- 1. Text extraction from vector/text layer
            text = page.get_text().strip()

            # --- 2. OCR extraction from image rendering
            pix = page.get_pixmap(dpi=300)
            img = Image.open(BytesIO(pix.tobytes("png"))).convert("RGB")
            ocr_text = pytesseract.image_to_string(img).strip()

            # --- 3. Combine both results
            page_text = f"\n--- Page {i} ---\n"
            if text:
                page_text += f"[PDF Text]\n{text}"
            if ocr_text:
                page_text += f"\n[OCR Text]\n{ocr_text}"

            # --- 4. Handle case when both are empty
            if not text and not ocr_text:
                page_text += "[No text found]"

            text_by_page.append(page_text)

    return "\n".join(text_by_page).strip()


def main():
    try:
        input_data = json.loads(sys.stdin.read())
        file_content = base64.b64decode(input_data["buffer"])
        mime_type = input_data["mimeType"]

        if mime_type == "application/pdf":
            result = extract_pdf_pagewise(file_content)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            result = extract_docx_text(file_content)
        elif mime_type.startswith("image"):
            result = extract_image_text(file_content)
        elif mime_type == "text/plain":
            result = extract_txt_text(file_content)
        elif "spreadsheet" in mime_type or "excel" in mime_type:
            result = extract_excel_text(file_content)
        else:
            result = "[Unsupported file type]"

        # ✅ Only send valid JSON to stdout
        print(json.dumps({"text": result}))

    except Exception as e:
        print(json.dumps({"error": str(e)}))


if __name__ == "__main__":
    main()

